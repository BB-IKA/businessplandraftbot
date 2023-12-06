import streamlit as st
import openai
from docx import Document

# Constants for sections of the business plan
SECTIONS_TEMPLATE = [
    "Executive Summary",
    "Company Description",
    "Market Analysis",
    "Organization & Management",
    "Service or Product Line",
    "Marketing & Sales",
    "Funding Request",
    "Financial Projections",
    "Appendix"
]

def main():
    st.title('AI Business Plan Generator 👨🏾‍💻')
    st.subheader('By: Better Digital Solutions')
    
    # Allow users to provide their OpenAI API key and select the model
    openai_api_key, openai_model = user_api_key_model_input()

    # Check if API key and model are provided
    if openai_api_key and openai_model:
        openai.api_key = openai_api_key

        # Initialize session state
        init_session_state()
        system_prompt = create_system_prompt()

        # UI for input fields
        input_ui()

        # Section content generation and navigation
        handle_section_generation(system_prompt, openai_model)

        # Business plan completion and download
        completion_ui()
    else:
        st.warning("Please enter your OpenAI API Key and select a model to start.")

def user_api_key_model_input():
    with st.sidebar:
        st.subheader("API Key and Model Configuration")
        api_key = st.text_input("Enter OpenAI API Key", type="password")
        model = st.selectbox("Select the Model", ["gpt-3.5-turbo", "text-davinci-003", "Other Models..."])
        return api_key, model

def init_session_state():
    if 'step' not in st.session_state:
        st.session_state.update({
            'step': -1,
            'business_name': '',
            'industry': '',
            'core_product': '',
            'problem_solved': '',
            'additional_info': '',
            'doc': Document(),
            'section_contents': [''] * len(SECTIONS_TEMPLATE),
        })
        st.session_state.doc.add_heading('Business Plan', level=0)

def input_ui():
    st.subheader("Enter details about your business:")
    with st.form(key='business_details_form'):
        st.session_state.business_name = st.text_input("Business Name", st.session_state.business_name, help="Enter the name of your business.")
        st.session_state.industry = st.text_input("Industry", st.session_state.industry, help="Enter the industry your business operates in.")
        st.session_state.core_product = st.text_input("Core Product/Service", st.session_state.core_product, help="Describe the core product or service your business offers.")
        st.session_state.problem_solved = st.text_input("Problem Solved", st.session_state.problem_solved, help="Explain the problem that your product or service solves.")
        st.session_state.additional_info = st.text_area("Additional Information", st.session_state.additional_info, help="Provide any additional information about your business.")
        submitted = st.form_submit_button("Save Business Details")
        if submitted:
            st.success("Business details saved!")

def handle_section_generation(system_prompt, model):
    if st.session_state.step == -1:
        if st.button("Start Generating"):
            st.session_state.step = 0
            st.session_state.section_contents[0] = generate_section_content(SECTIONS_TEMPLATE[0], system_prompt, model)
            st.rerun()
    else:
        section_content_ui(system_prompt, model)

def section_content_ui(system_prompt, model):
    if st.session_state.step >= 0:
        progress_percent = calculate_progress()
        st.progress(progress_percent)
        st.caption(f'Section {st.session_state.step + 1} of {len(SECTIONS_TEMPLATE)}')

        section = SECTIONS_TEMPLATE[st.session_state.step]
        with st.expander(f"{section}", expanded=True):
            st.text_area("Content", value=st.session_state.section_contents[st.session_state.step], height=200, key=f'{section}_text')

        navigation_buttons(system_prompt, model)

def navigation_buttons(system_prompt, model):
    col1, col2 = st.columns(2)
    with col1:
        if st.session_state.step > 0:
            if st.button("Previous Section"):
                st.session_state.step -= 1
    with col2:
        if st.session_state.step < len(SECTIONS_TEMPLATE) - 1:
            if st.button("Next Section"):
                st.session_state.step += 1
                st.session_state.section_contents[st.session_state.step] = generate_section_content(SECTIONS_TEMPLATE[st.session_state.step], system_prompt, model)
                st.rerun()
        else:
            if st.button("Complete"):
                st.session_state.step = len(SECTIONS_TEMPLATE)

def generate_section_content(section, system_prompt, model):
    try:
        user_input = (
            f"Elaborate on the '{section}' section for {st.session_state.business_name}, "
            f"a {st.session_state.industry} business offering {st.session_state.core_product} "
            f"to solve {st.session_state.problem_solved}. "
            f"Additional Info: {st.session_state.additional_info}"
        )
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_input}
            ]
        )
        content = response.choices[0]['message']['content']
        add_section_to_document(section, content)
        return content
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return ""

def add_section_to_document(section, content):
    st.session_state.doc.add_heading(section, level=1)
    st.session_state.doc.add_paragraph(content)

def completion_ui():
    if st.session_state.step >= len(SECTIONS_TEMPLATE):
        st.success("Business plan is complete!")
        save_path = "business_plan.docx"
        
        # Save the Word document
        st.session_state.doc.save(save_path)
        
        # Provide download button
        with open(save_path, "rb") as f:
            st.download_button("Download Business Plan", f, file_name="business_plan.docx")
        
        # Reset button
        if st.button("Start Over"):
            st.session_state.step = -1  # Reset to start

def calculate_progress():
    step = max(0, st.session_state.step)
    return step / (len(SECTIONS_TEMPLATE) - 1)

def create_system_prompt():
    return (
        "Generate a comprehensive business plan.\n"
        "The user is planning to start a business and needs a structured plan.\n"
        "1. Understand the business details provided by the user.\n"
        "2. Ensure the content is organized and coherent.\n"
        "3. Stay within the scope of a business plan.\n"
        "Based on the provided details, create a detailed section of a business plan.\n"
        "Business name, industry, core product/service, problem solved, and any additional information."
    )

if __name__ == "__main__":
    main()
